from django.shortcuts import render, redirect,get_object_or_404
from django.contrib.auth import authenticate, login
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.contrib.auth.models import User
from django.core.exceptions import ValidationError
from django.contrib.auth import logout
from .models import *
from django.utils.timezone import now
from django.core.paginator import Paginator
from django.db.models import F,DecimalField,ExpressionWrapper ,Value as V
from decimal import Decimal
from io import BytesIO
import pdfkit
from docx import Document
import os
from django.http import HttpResponse
from django.conf import settings
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from django.core.files.storage import default_storage
from docx.shared import Inches
from PIL import Image
from decimal import Decimal, InvalidOperation
from django.db.models import Sum
import locale
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import io
import base64
from django.db.models.functions import TruncMonth, TruncDay
import calendar
import json
from django.utils.dateformat import DateFormat
from django.shortcuts import render
from babel.dates import format_date

# Create your views here.
def agregar_emisor(request):
    if request.method == 'POST':
        dni = request.POST.get('dni')
        nombre = request.POST.get('nombre')
        apellido = request.POST.get('apellido')
        celular = request.POST.get('celular')
        imagen_dni1 = request.FILES.get('imagen_dni1')
        imagen_dni2 = request.FILES.get('imagen_dni2')

        Emisor.objects.create(
            dni=dni,
            nombre=nombre,
            apellido=apellido,
            celular=celular,
            imagen_dni1=imagen_dni1,
            imagen_dni2=imagen_dni2
        )
        messages.success(request, 'Emisor agregado')

        return redirect('listar_emisores')

    return render(request, 'agregar_emisor.html')

def declaracion_jurada(request):
    if request.method == 'POST':
        cliente_id = request.POST.get('cliente_id')
        emisor_id = request.POST.get('emisor_id')
        fecha=request.POST.get('fecha')
        cliente = Cliente.objects.get(id=cliente_id)
        emisor = Emisor.objects.get(id=emisor_id)
        tipo = request.POST.get('tipo')  # Nuevo campo para tipo

        # Ruta del archivo de plantilla en STATICFILES_DIRS
        template_path = os.path.join(settings.BASE_DIR, 'static', 'declaracion', 'declaracion.docx')

        # Cargar el documento base
        try:
            doc = Document(template_path)
        except Exception as e:
            messages.error(request, f'Error al cargar el documento: {e}')
            return redirect('home')

        # Reemplazar los placeholders con los datos reales
        nombre_completo_emisor = f"{emisor.nombre} {emisor.apellido}"
        nombre_completo_cliente = f"{cliente.nombre} {cliente.apellido}"
        placeholders = {
            'NOMBRE_EMISOR': nombre_completo_emisor,
            'DNI_EMISOR': emisor.dni,
            'NOMBRE_CLIENTE': nombre_completo_cliente,
            'DNI_CLIENTE': cliente.dni,
            'TELEFONO_EMISOR': emisor.celular,
            'TELEFONO_CLIENTE': cliente.celular,
            'SEDE':cliente.destino,
            'FECHA': fecha,
            'TIPO': tipo  # Reemplazar el nuevo placeholder
        }

        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in placeholders.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)

        # Reemplazar FOTO1 y FOTO2 con las imágenes correspondientes
        for paragraph in doc.paragraphs:
            if 'FOTO1' in paragraph.text:
                paragraph.clear()  # Limpiar el contenido de este párrafo
                run = paragraph.add_run()
                if emisor.imagen_dni1:
                    image_path1 = os.path.join(settings.MEDIA_ROOT, emisor.imagen_dni1.name)
                    run.add_picture(image_path1, width=Inches(4.1), height=Inches(2.75))
            if 'FOTO2' in paragraph.text:
                paragraph.clear()  # Limpiar el contenido de este párrafo
                run = paragraph.add_run()
                if emisor.imagen_dni2:
                    image_path2 = os.path.join(settings.MEDIA_ROOT, emisor.imagen_dni2.name)
                    run.add_picture(image_path2, width=Inches(4.1), height=Inches(2.75))

        # Guardar el documento temporalmente
        temp_doc_buffer = BytesIO()
        doc.save(temp_doc_buffer)
        temp_doc_buffer.seek(0)

        # Enviar el archivo Word como respuesta de descarga
        response = HttpResponse(temp_doc_buffer.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="declaracion_jurada_{cliente.dni}.docx"'

        # Mostrar mensaje de éxito
        return response

    return redirect('home')

def user_login(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')  # Redirige a una vista de inicio después del login
        else:
            messages.error(request, 'Nombre de usuario o contraseña incorrectos.')
    return render(request, 'login.html')

def listar_emisores(request):
    emisores_list = Emisor.objects.all()

    return render(request, 'listar_emisores.html', {'emisores_list': emisores_list })


from .forms import EmisorForm  # Importa el formulario que acabas de crear

def modificar_emisor(request, id):
    emisor = get_object_or_404(Emisor, id=id)

    if request.method == 'POST':
        form = EmisorForm(request.POST, request.FILES, instance=emisor)
        if form.is_valid():
            form.save()
            messages.success(request, 'Emisor modificado correctamente.')
            return redirect('listar_emisores')  # Redirige a la lista de emisores
    else:
        form = EmisorForm(instance=emisor)

    return render(request, 'modificar_emisor.html', {'form': form})

def eliminar_emisor(request, id):
    emisor = get_object_or_404(Emisor, id=id)
    emisor.eliminado = True  # Marca el emisor como eliminado
    emisor.save()
    messages.success(request, f"El emisor con ID {id} ha sido eliminado exitosamente.")
    return redirect('ver_emisores')  # Redirige a la lista de emisores
def eliminar_gasto(request, id):
    gasto = get_object_or_404(Gasto, id=id)
    gasto.eliminado = True  # Marca el gasto como eliminado
    gasto.save()
    messages.success(request, f"El gasto con ID {id} ha sido eliminado exitosamente.")
    return redirect('ver_gastos')  # Redirige a la lista de gastos

from .forms import GastoForm

def modificar_gasto(request, id):
    gasto = get_object_or_404(Gasto, id=id)

    if request.method == 'POST':
        form = GastoForm(request.POST, instance=gasto)
        if form.is_valid():
            form.save()
            messages.success(request, 'Gasto modificado correctamente.')

            return redirect('ver_gastos')
    else:
        form = GastoForm(instance=gasto)

    return render(request, 'modificar_gasto.html', {'form': form})
def user_register(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')

        # Validar las contraseñas
        if password1 != password2:
            messages.error(request, 'Las contraseñas no coinciden.')
            return redirect('register')

        # Crear el usuario
        try:
            user = User.objects.create_user(username=username, password=password1)
            messages.success(request, 'Registro exitoso')
            return redirect('login')
        except ValidationError as e:
            messages.error(request, f'Error al crear el usuario: {e.message}')
            return redirect('register')
    return render(request, 'register.html')
from django.db.models.functions import Coalesce
from django.db.models import Sum, F, Value, Case, When ,DecimalField

@login_required(login_url='/login/')
def home(request):
    year = request.GET.get('year')
    month = request.GET.get('month')
    selected_month = request.GET.get('month', datetime.now().month)

    if not year:
        year = datetime.now().year
    else:
        year = int(year)

    if not month:
        month = 0
    else:
        month = int(month)

    # Obtén el capital total
    capital = Capital.objects.first() or Capital.objects.create()

    # Calcula el capital en ropa disponible
    Capital_en_Ropa = Producto.objects.filter(estado='Disponible').aggregate(total=Sum('precio_compra'))['total'] or Decimal('0.00')

    # Filtrar ventas pendientes
    ventas_pendientes = Venta.objects.filter(estado='PENDIENTE')

    # Inicializar la variable para acumular el total a restar
    total_a_restar = Decimal('0.00')

    # Iterar sobre las ventas pendientes y calcular el total a restar
    for venta in ventas_pendientes:
        # Sumar todos los pagos fraccionados de la venta actual
        total_pagado = PagoFraccionado.objects.filter(venta=venta).aggregate(total=Sum('monto_pagado'))['total'] or Decimal('0.00')

        # Obtener el precio de compra del producto asociado a la venta
        precio_compra_producto = venta.producto.precio_compra

        # Si el total pagado es menor que el precio de compra del producto, calcular la diferencia
        if total_pagado < precio_compra_producto:
            diferencia = precio_compra_producto - total_pagado
            total_a_restar += diferencia

    # Calcular el saldo para invertir
    saldo_para_invertir = capital.monto - Capital_en_Ropa - total_a_restar

    today = timezone.now().date()
    ventas_del_dia_sin_fraccionado = Venta.objects.filter(
        fecha_compra__date=today,
        estado='COMPLETADO'
    ).exclude(pagos__isnull=False)
    ventas_del_dia = Venta.objects.filter(fecha_compra__date=today, estado='COMPLETADO')

    total_ventas_dia = ventas_del_dia_sin_fraccionado.aggregate(
        total=Sum('total_venta')
    )['total'] or Decimal('0.00')


    costo_ventas_dia= ventas_del_dia_sin_fraccionado.aggregate(
        total=Sum('producto__precio_compra')
    )['total'] or Decimal('0.00')
    # Ventas pendientes del día que tienen pagos fraccionados
    pagos_fraccionados_hoy = PagoFraccionado.objects.filter(fecha_pago__date=today)
    ganancia_pagos_fraccionados = pagos_fraccionados_hoy.aggregate(total=Sum('ganancia'))['total'] or Decimal('0.00')
    # Gastos del día
    gastos_del_dia = Gasto.objects.filter(fecha=today,eliminado=False)
    total_gastos_dia = gastos_del_dia.aggregate(total=Sum('monto'))['total'] or Decimal('0.00')
    # Calcular ganancia del día
    ganancia_dia1 = total_ventas_dia - costo_ventas_dia - total_gastos_dia + ganancia_pagos_fraccionados

    # Data for charts
    if month:
        start_date = timezone.now().replace(year=year, month=month, day=1)
        end_date = (start_date + timedelta(days=31)).replace(day=1)
        ventas = Venta.objects.filter(
            fecha_compra__range=[start_date, end_date],
            estado='COMPLETADO'
        ).annotate(day=TruncDay('fecha_compra')).values('day').annotate(
            total_ventas=Sum('total_venta'),
            costo_total=Sum('producto__precio_compra')
        ).order_by('day')

        gastos = Gasto.objects.filter(
            fecha__range=[start_date, end_date],
            eliminado=False
        ).annotate(day=TruncDay('fecha')).values('day').annotate(
            total_gastos=Sum('monto')
        ).order_by('day')

        dias = []
        ganancias_diarias = []

        for i in range(1, 32):
            try:
                dia = timezone.now().replace(year=year, month=month, day=i).strftime('%d-%m-%Y')
                ventas_dia = ventas.filter(day__day=i).first()
                gastos_dia = gastos.filter(day__day=i).first()
            except:
                continue

            total_ventas = ventas_dia['total_ventas'] if ventas_dia else Decimal('0.00')
            costo_total = ventas_dia['costo_total'] if ventas_dia else Decimal('0.00')
            total_gastos = gastos_dia['total_gastos'] if gastos_dia else Decimal('0.00')

            ganancia_dia = total_ventas - costo_total - total_gastos
            dias.append(dia)
            ganancias_diarias.append(float(ganancia_dia))

        context = {
            'capital': capital,
            'historial_movimientos': capital.historiales.all(),
            'Capital_en_Ropa': float(Capital_en_Ropa),
            'saldo_para_invertir': float(saldo_para_invertir),
            'ganancia_dia': float(ganancia_dia1),
            'dias': json.dumps(dias),
            'ganancias_diarias': json.dumps(ganancias_diarias),
            'selected_year': year,
            'selected_month': month,
        }
    else:
        ventas = Venta.objects.filter(
            fecha_compra__year=year,
            estado='COMPLETADO'
        ).annotate(month=TruncMonth('fecha_compra')).values('month').annotate(
            total_ventas=Sum('total_venta'),
            costo_total=Sum('producto__precio_compra')
        ).order_by('month')

        gastos = Gasto.objects.filter(
            fecha__year=year,
            eliminado=False  # Filtrar solo gastos que no están eliminados

        ).annotate(month=TruncMonth('fecha')).values('month').annotate(
            total_gastos=Sum('monto')
        ).order_by('month')

        meses = []
        ganancias_mensuales = []
        # Calcular el saldo pendiente y el saldo para invertir (YY) para cada venta pendiente
        saldos_pendientes = Venta.objects.filter(estado='PENDIENTE').annotate(
            total_pagado=Coalesce(Sum('pagos__monto_pagado'), V(0), output_field=DecimalField()),
            precio_compra_producto=F('producto__precio_compra'),
            saldo_para_invertir=F('precio_compra_producto') - F('total_pagado'),
            ganancia=F('total_venta') - F('precio_compra_producto')
        ).values('cliente__nombre', 'total_pagado', 'precio_compra_producto', 'saldo_para_invertir', 'ganancia')

        # Formatear los saldos pendientes en el formato deseado
        saldos_pendientes_formateados = [
            {
                'cliente': item['cliente__nombre'],
                'saldo_formateado': f"${float(max(Decimal('0.00'), item['saldo_para_invertir'])):.2f} ({float(item['total_pagado']):.2f} para invertir / {float(item['ganancia']):.2f} ganancia)"
            }
            for item in saldos_pendientes
            if item['total_pagado'] < item['precio_compra_producto']
        ]



        for i in range(1, 13):
            mes = timezone.now().replace(year=year, month=i, day=1).strftime('%B %Y')
            ventas_mes = ventas.filter(month__month=i).first()
            gastos_mes = gastos.filter(month__month=i).first()

            total_ventas = ventas_mes['total_ventas'] if ventas_mes else Decimal('0.00')
            costo_total = ventas_mes['costo_total'] if ventas_mes else Decimal('0.00')
            total_gastos = gastos_mes['total_gastos'] if gastos_mes else Decimal('0.00')

            ganancia_mes = total_ventas - costo_total - total_gastos
            meses.append(mes)
            ganancias_mensuales.append(float(ganancia_mes))

        context = {
            'capital': capital,
            'historial_movimientos': capital.historiales.all(),
            'Capital_en_Ropa': float(Capital_en_Ropa),
            'saldo_para_invertir': float(saldo_para_invertir),
            'ganancia_dia': float(ganancia_dia1),
            'meses': json.dumps(meses),
            'ganancias_mensuales': json.dumps(ganancias_mensuales),
            'selected_year': year,
            'selected_month': month,
            'saldos_pendientes' : saldos_pendientes_formateados,
        }

    return render(request, 'home.html', context)


@login_required(login_url='/login/')
def aumentar_capital(request):
    if request.method == 'POST':
        monto_aumentar = request.POST.get('monto_aumentar', '0')

        try:
            monto_aumentar = Decimal(monto_aumentar)
            if monto_aumentar > 0:
                capital, created = Capital.objects.get_or_create(pk=1)  # Asumiendo que solo hay una instancia de Capital

                capital.aumentar_capital(monto_aumentar)

                messages.success(request, f'Se ha aumentado el capital en {monto_aumentar:.2f}')
            else:
                messages.error(request, 'El monto a aumentar debe ser mayor a 0')
        except (ValueError, InvalidOperation):
            messages.error(request, 'Monto inválido')
    return redirect('home')

def ver_clientes(request):
    clientes_list = Cliente.objects.all()
    emisores = Emisor.objects.all()  # Obtener todos los emisores
    paginator = Paginator(clientes_list, 8)  # Mostrar 8 clientes por página
    page_number = request.GET.get('page')
    clientes = paginator.get_page(page_number)

    return render(request, 'ver_clientes.html', {'clientes': clientes, 'emisores': emisores})
def ver_ventas(request):
    estado_filtro = request.GET.get('estado', 'TODOS')

    if estado_filtro == 'TODOS':
        ventas_list = Venta.objects.all().order_by('-fecha_compra')
    else:
        ventas_list = Venta.objects.filter(estado=estado_filtro).order_by('-fecha_compra')

    # Calcular la ganancia para cada venta
    ventas_con_ganancia = []
    for venta in ventas_list:
        producto = venta.producto
        ganancia = venta.total_venta - producto.precio_compra
        fecha_formateada = DateFormat(venta.fecha_compra).format('Y-m-d')
        ventas_con_ganancia.append({
            'venta': venta,
            'ganancia': ganancia,
            'fecha_formateada': fecha_formateada
        })

    paginator = Paginator(ventas_con_ganancia, 8)  # Mostrar 8 ventas por página
    page_number = request.GET.get('page', 1)  # Obtener el número de página desde la solicitud GET
    page_obj = paginator.get_page(page_number)

    return render(request, 'ventas.html', {
        'page_obj': page_obj,
        'estado_filtro': estado_filtro
    })






def agregar_pago(request):
    if request.method == 'POST':
        venta_id = request.POST.get('venta')
        nuevo_pago_str = request.POST.get('nuevo_pago')

        if not venta_id or not nuevo_pago_str:
            messages.error(request, "Faltan datos para procesar el pago.")
            return redirect('agregar_pago')

        try:
            venta = get_object_or_404(Venta, id=venta_id)
            nuevo_pago = Decimal(nuevo_pago_str)  # Convierte el string a Decimal

            if venta.pago_fraccionado is None:
                venta.pago_fraccionado = Decimal('0')  # Asegúrate de que sea Decimal

            venta.pago_fraccionado += nuevo_pago

            if venta.pago_fraccionado >= venta.total_venta:
                venta.estado = 'COMPLETADO'

            venta.save()
            messages.success(request, 'Pago agregado exitosamente.')

        except ValueError:
            messages.error(request, "El monto del nuevo pago no es válido.")
        except Exception as e:
            messages.error(request, f"Error inesperado: {e}")

        return redirect('ver_ventas')

    else:
        ventas_fraccionadas = Venta.objects.filter(
            pago_fraccionado__lt=F('total_venta'),
            estado='PENDIENTE'
        ).order_by('-fecha_compra')
        return render(request, 'agregar_pago.html', {'ventas_fraccionadas': ventas_fraccionadas})@login_required(login_url='/login/')


def agregar_pago(request):
    if request.method == 'POST':
        venta_id = request.POST.get('venta')
        nuevo_pago_str = request.POST.get('nuevo_pago')

        if not venta_id or not nuevo_pago_str:
            messages.error(request, "Faltan datos para procesar el pago.")
            return redirect('agregar_pago')

        try:
            venta = get_object_or_404(Venta, id=venta_id)
            nuevo_pago = Decimal(nuevo_pago_str)  # Convierte el string a Decimal

            # Verifica si el pago fraccionado es nulo o no existe
            if venta.pago_fraccionado is None:
                venta.pago_fraccionado = Decimal('0')  # Asegúrate de que sea Decimal

            # Calcula el nuevo monto total pagado después del pago actual
            nuevo_monto_pagado = venta.pago_fraccionado + nuevo_pago

            # Calcula la ganancia
            ganancia = max(nuevo_pago - max(venta.producto.precio_compra - venta.pago_fraccionado, 0), Decimal('0'))


            # Crear una nueva instancia de PagoFraccionado
            PagoFraccionado.objects.create(
                venta=venta,
                monto_pagado=nuevo_pago,
                fecha_pago=timezone.now(),
                ganancia=ganancia
            )
            # Actualiza el total pagado en la venta
            venta.pago_fraccionado = nuevo_monto_pagado

            # Cambia el estado de la venta si el total pagado cubre o excede el total de la venta
            if venta.pago_fraccionado >= venta.total_venta:
                venta.estado = 'COMPLETADO'

            venta.save()
            messages.success(request, 'Pago agregado exitosamente.')

        except ValueError:
            messages.error(request, "El monto del nuevo pago no es válido.")
        except Exception as e:
            messages.error(request, f"Error inesperado: {e}")

        return redirect('ver_ventas')

    else:
        ventas_fraccionadas = Venta.objects.filter(
            pago_fraccionado__lt=F('total_venta'),
            estado='PENDIENTE'
        ).order_by('-fecha_compra')
        return render(request, 'agregar_pago.html', {'ventas_fraccionadas': ventas_fraccionadas})


def ver_gastos(request):
    gastos_list = Gasto.objects.filter(eliminado=False)  # Filtra para mostrar solo los gastos no eliminados
    gastos_formateados = []

    for gasto in gastos_list:
        fecha_formateada = format_date(gasto.fecha, format='d MMMM yyyy', locale='es')
        gastos_formateados.append({
            'id': gasto.id,  # Agregamos el ID aquí
            'tipo_gasto' : gasto.get_tipo_gasto_display(),
            'fecha': fecha_formateada,
            'descripcion': gasto.descripcion,
            'monto': gasto.monto,
        })

    return render(request, 'ver_gastos.html', {'gastos': gastos_formateados})




def crear_venta(request):
    if request.method == 'POST':
        producto_nombre = request.POST.get('producto')  # El nombre del producto desde el input autocompletado
        cliente_id = request.POST.get('cliente')
        precio_venta = request.POST.get('precio_venta')
        tipo_venta = request.POST.get('tipo_venta')
        pago_fraccionado = request.POST.get('pago_fraccionado', None)  # Puede ser None si no se proporciona

        try:
            producto = Producto.objects.get(nombre=producto_nombre)
            cliente = Cliente.objects.get(id=cliente_id)

            # Determinar el estado de la venta basado en el tipo de venta
            estado_venta = 'PENDIENTE' if tipo_venta == 'FRACCIONADO' else 'COMPLETADO'

            # Crear la venta
            venta = Venta(
                producto=producto,
                cliente=cliente,
                tipo_venta=tipo_venta,
                estado=estado_venta,
                total_venta=precio_venta,  # Guardar el precio de venta como total_venta
                fecha_compra=timezone.now()  # Asignar la fecha de la venta a la fecha y hora actual

            )
            venta.save()

            # Actualizar el precio de venta del producto y cambiar su estado a 'Vendido'
            producto.precio_venta = precio_venta
            producto.estado = 'Vendido'
            producto.save()

            if tipo_venta == 'FRACCIONADO' and pago_fraccionado:
                # Crear la instancia de PagoFraccionado
                pago = PagoFraccionado(
                    venta=venta,
                    monto_pagado=Decimal(pago_fraccionado),
                    fecha_pago=timezone.now()
                )
                # Calcular la ganancia
                pago.ganancia = max(pago.monto_pagado - producto.precio_compra, 0)
                pago.save()

                # Actualizar el monto total de pago fraccionado en la venta
                venta.pago_fraccionado = pago_fraccionado
                venta.save()

            messages.success(request, f'Venta creada con éxito. Producto: {producto.nombre}, Cliente: {cliente.nombre} {cliente.apellido}.')
            return redirect('ver_ventas')

        except Producto.DoesNotExist:
            messages.error(request, 'Producto no encontrado.')
            return redirect('crear_venta')

        except Cliente.DoesNotExist:
            messages.error(request, 'Cliente no encontrado.')
            return redirect('crear_venta')

    else:
        productos = Producto.objects.filter(estado='Disponible')
        clientes = Cliente.objects.all()
        return render(request, 'crear_venta.html', {'productos': productos, 'clientes': clientes})






def ver_dropshipping(request):
    # Lógica para mostrar dropshipping
    return render(request, 'ver_dropshipping.html')

def ver_productos(request):
    estado = request.GET.get('estado', None)
    if estado:
        productos_list = Producto.objects.filter(estado=estado).order_by('-fecha_subida')
    else:
        productos_list = Producto.objects.all().order_by('-fecha_subida')

    paginator = Paginator(productos_list, 8)  # Mostrar 8 productos por página
    page_number = request.GET.get('page')
    productos = paginator.get_page(page_number)

    return render(request, 'ver_productos.html', {'productos': productos})
def salir_view(request):
    # Cerrar la sesión del usuario
    logout(request)
    # Redirigir al inicio de sesión
    return redirect('login')

def perfil_view(request):
    context = {
        'user': request.user
    }
    return render(request, 'perfil.html', context)
def crear_producto(request):
    if request.method == 'POST':
        # Obtener datos del formulario
        nombre = request.POST.get('nombre')
        talla = request.POST.get('talla')
        marca = request.POST.get('marca')
        descripcion = request.POST.get('descripcion')
        precio_compra = Decimal(request.POST.get('precio_compra'))
        precio_venta = request.POST.get('precio_venta')
        fecha_compra = request.POST.get('fecha_compra')
        categoria = request.POST.get('categoria')

        # Validar los datos (puedes agregar más validaciones si es necesario)
        if not all([nombre, talla, marca, descripcion, precio_compra, precio_venta, fecha_compra, categoria]):
            messages.error(request, 'Todos los campos son obligatorios.')
            return render(request, 'producto.html')

        # Obtener el capital y calcular el saldo para invertir
        capital = Capital.objects.first()  # Asumiendo que solo hay un registro de capital
        if not capital:
            capital = Capital.objects.create()  # Crear el registro si no existe

        # Calcular el saldo invertido total
        Capital_en_Ropa = Producto.objects.filter(estado='Disponible').aggregate(total=Sum('precio_compra'))['total'] or Decimal('0.00')

        # Calcular el saldo para invertir
        saldo_para_invertir = capital.monto - Capital_en_Ropa

        # Validar si hay suficiente saldo para la compra
        if saldo_para_invertir < precio_compra:
            faltante = precio_compra - saldo_para_invertir
            messages.error(request, f'Fondos insuficientes para agregar el producto. Te faltan {faltante} para completar la compra.')
            return redirect('ver_productos')

        try:
            # Crear una nueva instancia de Producto
            producto = Producto(
                nombre=nombre,
                talla=talla,
                marca=marca,
                descripcion=descripcion,
                precio_compra=precio_compra,
                precio_venta=precio_venta,
                fecha_compra=fecha_compra,
                categoria=categoria,
            )
            producto.save()
            messages.success(request, f'Producto "{nombre}" agregado correctamente.')
            return redirect('ver_productos')  # Redirige a la página de productos o a donde desees
        except Exception as e:
            messages.error(request, f'Error al agregar el producto: {e}')
            return render(request, 'producto.html')

    return render(request, 'producto.html')



def crear_gasto(request):
    if request.method == 'POST':
        tipo_gasto = request.POST.get('tipo_gasto')
        monto = request.POST.get('monto')
        fecha = request.POST.get('fecha')
        descripcion = request.POST.get('descripcion')

        if tipo_gasto and monto and fecha:
            Gasto.objects.create(
                tipo_gasto=tipo_gasto,
                monto=monto,
                fecha=fecha,
                descripcion=descripcion
            )
            messages.success(request, 'Gasto agregado correctamente.')

            return redirect('ver_gastos')

    # Render the form without passing gastos
    return render(request, 'gastos.html')



@login_required(login_url='/login/')
def ver_reportes(request):
    # Obtener mes y año seleccionados por el usuario o usar el actual
    mes_actual = int(request.GET.get('mes', now().month))
    año_actual = int(request.GET.get('año', now().year))

    # Agrupar ventas por día del mes seleccionado y calcular la ganancia diaria
    ventas = Venta.objects.filter(fecha_compra__month=mes_actual, fecha_compra__year=año_actual, estado='COMPLETADO') \
        .annotate(dia=F('fecha_compra__day')) \
        .values('dia') \
        .annotate(ganancia_dia=Sum(ExpressionWrapper(F('total_venta') - F('producto__precio_compra'), output_field=DecimalField())))

    # Agrupar gastos por día del mes seleccionado
    gastos = Gasto.objects.filter(fecha__month=mes_actual, fecha__year=año_actual,eliminado=False) \
        .values('fecha') \
        .annotate(total_gasto=Sum('monto'))

    # Combinar ventas y gastos por día para calcular la ganancia diaria neta
    reporte = []
    dias_procesados = set()

    for venta in ventas:
        dia = venta['dia']
        ganancia_dia = venta['ganancia_dia']
        gasto_dia = next((g['total_gasto'] for g in gastos if g['fecha'].day == dia), 0)
        ganancia_diaria = ganancia_dia - gasto_dia

        reporte.append({
            'fecha': f"{dia}-{mes_actual}-{año_actual}",
            'ganancia_dia': ganancia_dia,
            'gasto_dia': gasto_dia,
            'ganancia_diaria': ganancia_diaria
        })

        dias_procesados.add(dia)

    # Añadir días con solo gastos (sin ventas)
    for gasto in gastos:
        dia = gasto['fecha'].day
        if dia not in dias_procesados:
            reporte.append({
                'fecha': f"{dia}-{mes_actual}-{año_actual}",
                'ganancia_dia': 0,
                'gasto_dia': gasto['total_gasto'],
                'ganancia_diaria': -gasto['total_gasto']
            })

    reporte.sort(key=lambda x: int(x['fecha'].split('-')[0]))

    # Nombre del reporte
    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
    mes_nombre = datetime(year=año_actual, month=mes_actual, day=1).strftime('%B')
    nombre_reporte = f"Reporte de {mes_nombre.capitalize()} {año_actual}"

    context = {
        'reporte': reporte,
        'nombre_reporte': nombre_reporte,
        'mes_actual': mes_actual,
        'año_actual': año_actual,
    }

    return render(request, 'ver_reportes.html', context)



def crear_cliente(request):
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        apellido = request.POST.get('apellido')
        celular = request.POST.get('celular')
        dni = request.POST.get('dni')
        destino = request.POST.get('destino')

        if nombre and apellido and celular and dni and destino:
            Cliente.objects.create(
                nombre=nombre,
                apellido=apellido,
                celular=celular,
                dni=dni,
                destino=destino
            )
            messages.success(request, 'Cliente agregado correctamente.')
            return redirect('ver_clientes')

    return render(request, 'crear_cliente.html')

def ventas_view(request):
    return render(request, 'ventas.html')

def gastos_view(request):
    return render(request, 'gastos.html')